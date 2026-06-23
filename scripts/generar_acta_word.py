"""Genera docs/acta_reclamacion_40_puntos.docx — acta oficial SIDACS."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
CAPTURAS = ROOT / "docs" / "evidencias" / "capturas"
OUT = ROOT / "docs" / "acta_reclamacion_40_puntos.docx"
OUT_ALT = ROOT / "docs" / "acta_reclamacion_40_puntos_ENTREGA.docx"

# Tabla oficial de planificación (Sección 1)
SPRINTS = [
    ("Sprint 0", "22/05/2026", "Investigación inicial, mockups y arquitectura preliminar. Rol: Cliente final.",
     "Mockups y diseño inicial de interfaz", "Investigación técnica y apoyo en arquitectura lógica", "—"),
    ("Sprint 1", "28/05/2026", "Captura de video e importación funcional. Rol: Stakeholder técnico. Commits: c59daa5, 7c38538.",
     "Interfaz de carga de video y botones", "Captura de video y procesamiento inicial", "6"),
    ("Sprint 2", "01/06/2026", "Detección básica de movimiento mediante Farneback. Rol: Cliente de negocio. Commits: aa7a41e, 6fdd1c6.",
     "Visualización gráfica del movimiento", "Implementación básica de flujo óptico", "6"),
    ("Sprint 3", "08/06/2026", "Detección de comportamientos sospechosos. Rol: Dueño de la empresa. Commits: eb1abc6, 1e2c0d8, 2860c07.",
     "Sistema visual de alertas", "Detección de anomalías y análisis de movimiento (ML)", "7"),
    ("Sprint 4", "12/06/2026", "Dashboard completo y estadísticas visuales. Rol: Cliente final. Commits: 0b8a8ef, d2b84d0.",
     "Dashboard funcional e interfaz final", "Integración lógica y optimización", "7"),
    ("Sprint 5", "15/06/2026", "Pruebas generales y optimización básica. Rol: Stakeholder. Commits: 28aa355, 675bdd4.",
     "Correcciones visuales y ajustes de UI", "Corrección de errores y optimización del sistema", "7"),
    ("Sprint 6", "15/06/2026", "Sistema funcional, demo y documentación final. Rol: Cliente final. Commit: 741863c.",
     "Presentación visual y dashboard", "Explicación técnica y soporte en la demo final", "7"),
]

# Evidencias por sprint — capturas reales presentadas
EVIDENCIAS = [
    ("Sprint 0 — Planificación (22/05/2026)",
     "Investigación inicial, mockups y arquitectura preliminar.",
     "Joshua: mockups Figma. Marcelo: arquitectura lógica.",
     [("00_mockup_figma_sprint0.png", "Mockup de interfaz diseñado en Figma.")]),
    ("Sprint 1 — Captura e importación de video (28/05/2026)",
     "Captura de video e importación funcional.",
     "Joshua: interfaz y botones. Marcelo: procesamiento de video.",
     [
         ("01_sprint1_interfaz_inicial.png", "Interfaz principal con modelo ML cargado."),
         ("02_sprint1_importar_video.png", "Diálogo de importación desde videoPrueba/."),
         ("03_sprint1_video_cargado.png", "Video importado listo para monitoreo."),
     ]),
    ("Sprint 2 — Flujo óptico Farneback (01/06/2026)",
     "Detección básica de movimiento y visualización gráfica.",
     "Joshua: visualización. Marcelo: implementación Farneback.",
     [("04_sprint2_3_analisis_activo.png", "Análisis activo: Farneback, marcadores en timeline, log de actividad.")]),
    ("Sprint 3 — Detección de comportamientos sospechosos (08/06/2026)",
     "Alertas visuales, predicción ML y capturas automáticas.",
     "Joshua: alertas visuales. Marcelo: detección ML y anomalías.",
     [("04_sprint2_3_analisis_activo.png", "Predicción ML, alertas de movimiento y confianza en tiempo real.")]),
    ("Sprint 4 — Dashboard y estadísticas visuales (12/06/2026)",
     "Dashboard completo y centro de evidencias.",
     "Joshua: dashboard e interfaz. Marcelo: integración y optimización.",
     [
         ("05_sprint4_evidencias_lista.png", "Pestaña Evidencias: capturas, clips e incidentes."),
         ("06_sprint4_evidencias_detalle.png", "Detalle de captura sospechosa (confianza 88%)."),
     ]),
    ("Sprint 5 — Testing y optimización (15/06/2026)",
     "Pruebas generales, métricas y optimización del sistema.",
     "Joshua: ajustes UI. Marcelo: tests y optimización.",
     [
         ("07_sprint5_metricas.png", "Pestaña Métricas: actividad, movimiento, riesgo."),
         ("08_sprint5_metricas_modelo.png", "Accuracy 86,96% y rendimiento CPU/RAM."),
         ("09_sprint5_6_reporte_pruebas.png", "Reporte de pruebas automatizadas."),
     ]),
    ("Sprint 6 — Demo y documentación final (15/06/2026)",
     "Sistema funcional, demo y documentación consolidada.",
     "Joshua: presentación visual. Marcelo: soporte técnico demo.",
     []),
]

REUNIONES = [
    ("22/05/2026", "Sprint 0", "Planificación, mockups Figma", "Remoto", "Aprobación plan de sprints"),
    ("28/05/2026", "Sprint 1", "Importar video, webcam, UI base", "Remoto", "Captura funcional"),
    ("01/06/2026", "Sprint 2", "Flujo óptico Farneback", "Remoto", "Detección movimiento validada"),
    ("08/06/2026", "Sprint 3", "Alertas, ML, capturas sospechosas", "Remoto", "Demo anomalías"),
    ("12/06/2026", "Sprint 4", "Dashboard, centro de evidencias", "Remoto", "Interfaz final aprobada"),
    ("15/06/2026", "Sprint 5 y 6", "Testing, métricas, demo, documentación", "Remoto", "Cierre técnico"),
]

PUNTAJE = [
    ("Sprint 1", "Captura e importación de video funcional", "6", "7"),
    ("Sprint 2", "Flujo óptico Farneback con visualización", "6", "7"),
    ("Sprint 3", "Detección de anomalías y alertas", "7", "8"),
    ("Sprint 4", "Dashboard e interfaz gráfica final", "7", "8"),
    ("Sprint 5", "Testing, métricas y optimización", "7", "8"),
    ("Sprint 6", "Demo y documentación final", "7", "8"),
]


def _font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(run, size=16, bold=True)
    elif level == 1:
        _font(run, size=13, bold=True)
    else:
        _font(run, size=11, bold=True)


def _para(doc, text, center=False, bold=False, size=11, italic=False, color=None):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        sh = c._element.get_or_add_tcPr()
        sh.append(sh.makeelement(qn("w:shd"), {qn("w:fill"): "E8E8E8", qn("w:val"): "clear"}))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def _img(doc, fname, w=6.0):
    p = CAPTURAS / fname
    if not p.exists():
        _para(doc, f"[Imagen no encontrada: {fname}]", italic=True, size=9)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(p), width=Inches(w))


def main() -> None:
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(1.6)

    _heading(doc, "Acta de Reclamación de Pago", 0)
    for t in [
        "Sistema Inteligente de Detección y Análisis de Comportamientos Sospechosos en Tiempo Real mediante Flujo Óptico y Visión por Computadora (SIDACS)",
        "Puntos reclamados: 40 puntos (Sprints 1–6) · Integrantes: Joshua Jair Chavez Abirari · Marcelo Mena Molina",
        "Docente / cliente: _________________________________ · Fecha: 17 de junio de 2026",
        "Repositorio: github.com/Mystbeast1234263/SistemaDeteccion1 · Commit: 741863c (15/06/2026)",
    ]:
        _para(doc, t, center=True, size=10)
    _para(doc, "Plazo de reclamación: miércoles 17/06/2026, 23:59 h.", center=True, bold=True, color=(128, 0, 0))

    _heading(doc, "1. Acta de sprints — Fechas, tareas y responsables", 1)
    _para(doc, "Planificación (Sprint 0) y seis sprints de desarrollo completados conforme al documento presentado el 22/05/2026.", size=10)
    _table(doc, ["Sprint", "Fecha", "Entregable general", "Joshua", "Marcelo", "Pts."], SPRINTS)
    _para(doc, "* Sprint 6 planificado 18/06/2026; cierre técnico el 15/06/2026.", size=9, italic=True)

    _heading(doc, "2. Evidencias por sprint", 1)
    _para(doc, "Capturas reales de la aplicación (17/06/2026), agrupadas según lo presentado en cada sprint.", size=10)
    n = 1
    for titulo, entregable, responsables, imgs in EVIDENCIAS:
        _heading(doc, titulo, 2)
        _para(doc, f"Entregable: {entregable}", size=10, italic=True)
        _para(doc, responsables, size=10)
        if not imgs:
            _para(doc, "Documentación: documentacion.html, manual_tecnico.html. Repositorio Git commit 741863c.", size=10)
        for fname, desc in imgs:
            _img(doc, fname)
            _para(doc, f"Evidencia {n} — {desc}", italic=True, size=9)
            n += 1

    _heading(doc, "2.1 Repositorio y documentación", 2)
    _table(doc, ["Recurso", "Descripción"], [
        ("GitHub — rama main", "741863c1ce0828e0b2e48e4fbc50f9bc7116b4bb"),
        ("documentacion.html", "Manual completo Sprints 0–6"),
        ("manual_tecnico.html", "Arquitectura técnica"),
        ("testing/test_results.json", "Resultados pruebas Sprint 5"),
        ("models/modelo_v1.pkl", "Modelo ML entrenado"),
    ])

    _heading(doc, "3. Registro de reuniones con el docente", 1)
    _table(doc, ["Fecha", "Sprint", "Avance demostrado", "Modalidad", "Observaciones"], REUNIONES)

    _heading(doc, "4. Cálculo del puntaje — 40 puntos", 1)
    _table(doc, ["Sprint", "Criterio", "Autoeval.", "Máx."], PUNTAJE + [("TOTAL RECLAMADO", "", "40", "46")])
    _table(doc, ["Criterio empresa", "Pts.", "Justificación"], [
        ("Completitud", "10", "S0–S6 + Git + documentacion.html + capturas por sprint"),
        ("Evidencias tangibles", "10", "Capturas reales, repo, test_results.json, modelo .pkl"),
        ("Coherencia", "10", "40 pts = S1–S6; fechas y commits verificables"),
        ("Profesionalismo", "10", "Acta formal con tablas, firmas y evidencias"),
        ("Total", "40", ""),
    ])

    _heading(doc, "5. Declaración jurada bajo compromiso", 1)
    box = doc.add_table(1, 1)
    box.style = "Table Grid"
    cell = box.rows[0].cells[0]
    sh = cell._element.get_or_add_tcPr()
    sh.append(sh.makeelement(qn("w:shd"), {qn("w:fill"): "FAFAFA", qn("w:val"): "clear"}))
    lines = [
        "Nosotros, Joshua Jair Chavez Abirari y Marcelo Mena Molina, declaramos bajo fe de verdad que:",
        "1. Completamos Sprint 0 y los seis sprints de desarrollo descritos en esta acta.",
        "2. Presentamos avances al docente en las fechas de la Sección 3.",
        "3. El software en el repositorio GitHub cumple lo pactado en cada sprint.",
        "4. Las capturas, commits y documentos son auténticos.",
        "5. Reclamamos 40 (cuarenta) puntos conforme a la Sección 4.",
        "Plazo: miércoles 17 de junio de 2026, 23:59 h.",
    ]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _font(p.add_run(line), size=10, bold=(i == 0))
    doc.add_paragraph()

    _para(doc, "Firmas del equipo:", bold=True)
    sig = doc.add_table(1, 2)
    sig.style = "Table Grid"
    sig.rows[0].cells[0].paragraphs[0].add_run("\n\n\nJoshua Jair Chavez Abirari\nCI: ___________  Firma: ___________")
    sig.rows[0].cells[1].paragraphs[0].add_run("\n\n\nMarcelo Mena Molina\nCI: ___________  Firma: ___________")
    _para(doc, "\nAceptación docente / cliente:\n\n\nNombre, cargo, firma, fecha")
    _para(doc, "SIDACS v6.0.0-sprint6 — Acta generada 17/06/2026", center=True, size=9, italic=True, color=(85, 85, 85))

    for path in (OUT, OUT_ALT):
        try:
            doc.save(str(path))
            print(f"OK: {path} ({n - 1} evidencias)")
            return
        except PermissionError:
            continue
    print(f"ERROR: cierre el Word abierto. Guardado alternativo falló.")


if __name__ == "__main__":
    main()
